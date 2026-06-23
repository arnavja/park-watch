"""Build a readable Word doc of the demo-video script."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

OUT = Path(__file__).parent.parent / "outputs" / "Demo_Video_Script.docx"

NAVY = RGBColor(0x0B, 0x2A, 0x4A)
ORANGE = RGBColor(0xFF, 0x6B, 0x2B)
GREY = RGBColor(0x60, 0x60, 0x60)


def _heading(doc, text, size=22, color=NAVY, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def _section_header(doc, timestamp, action):
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{timestamp}]   ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = ORANGE
    r2 = p.add_run(action)
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def _script_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.name = "Georgia"
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(6)


def _note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(8)


def build():
    doc = Document()

    # Page margins
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # ── Title block
    _heading(doc, "Park-Watch — Demo Video Script (5 min)", size=22, after=2)
    _heading(doc, "Gridlock Hackathon · Theme 1 · Flipkart HQ × BTP",
             size=11, color=GREY, after=14)

    # ── Recording setup
    _heading(doc, "Before you hit record", size=14, after=4)
    for line in [
        "•  Open https://park-watch.streamlit.app in a fresh browser window — no other tabs visible.",
        "•  Quit other apps so the recording stays clean.",
        "•  Use QuickTime: File → New Screen Recording. Click the arrow next to Record → enable Built-in Microphone.",
        "•  Browser zoom 100%, 1080p output.",
        "•  Practice once before the real take. Speak slowly. Leave 1–2 seconds of silence at start and end.",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.15)

    doc.add_paragraph()  # spacer

    # ── Script sections
    _heading(doc, "The script", size=16, after=4)

    sections = [
        ("0:00 – 0:15",
         "Wait 2 seconds after hitting record. Dashboard header visible.",
         "Hi, I'm Arnav Jain, Team Byte_me_kaar. This is Park-Watch — Bengaluru parking "
         "enforcement intelligence, my submission for Theme 1 of the Gridlock Hackathon. "
         "BTP collects almost 300,000 violation records in 5 months. The problem isn't "
         "data — it's that BTP can't see the patterns inside it or act on them in real "
         "time. Park-Watch fixes both."),

        ("0:15 – 0:35",
         "Briefly hover over the dashboard so the BTP theme framing lands. Slow cursor.",
         "BTP's brief flags three things — enforcement is reactive, there's no heatmap "
         "of violations versus congestion, and zone priority is unclear. I built 4 "
         "analytical modules to address each, combined into one dashboard."),

        ("0:35 – 1:30",
         "Hover over the bar chart titled 'The 12-hour enforcement blind spot'.",
         "Before I explain the modules, look at this chart. These are the hours when BTP "
         "books violations. Tall bars all morning, peaking around 10 to 11 AM. After 2 PM, "
         "the chart falls off a cliff. By 3 PM, it's near zero. At 7 PM, across 5 entire "
         "months, only 27 violations were booked in all of Bengaluru. 27.\n\n"
         "That's not because illegal parking stops. That's because officers go home. BTP is "
         "structurally blind for roughly 12 hours every single day.\n\n"
         "This gap is the operational insight Park-Watch is built around. The rest of the "
         "system — cost quantification, forecasting, patrol routing — all exist to close it."),

        ("1:30 – 2:30",
         "Scroll down to the 'Congestion cost' section. Hover briefly over the ₹389 Cr metric.",
         "First question — what does this cost the city?\n\n"
         "For each of the 381 hotspots that my DBSCAN clustering identified, I snap to the "
         "nearest road segment using OpenStreetMap and pull the real lane count. Then I "
         "apply the BPR delay function — the standard equation used by every "
         "transport-planning authority in the world. When capacity drops because a lane is "
         "blocked by illegal parking, delay grows non-linearly.\n\n"
         "Computed across the top 100 hotspots — which together cover over 90 percent of "
         "all violations — the city loses ₹389 crore a year. That's roughly 10 percent of "
         "Bengaluru's total congestion bill, attributable to just 100 parking zones.\n\n"
         "The single biggest one — KR Market Junction, in Upparpet — costs ₹16 crore per "
         "month, on its own. And every assumption in this calculation — dwell time, value "
         "of time, lane capacity — is an explicit, defensible knob."),

        ("2:30 – 3:35",
         "Scroll to 'Blind-spot forecast'. Hover over the R² metric so the train→test gap delta is visible.",
         "Next, I built a forecasting model that predicts where violations will happen "
         "during the blind spot.\n\n"
         "XGBoost regression on 1.38 million hourly observations. The features cover "
         "calendar effects, recent activity lags, geographic location, vehicle mix — and "
         "an annual seasonality layer that learns festival, monsoon, and holiday patterns. "
         "Diwali week is structurally different from a normal week in KR Market — the "
         "model picks that up. Recent data is weighted higher via a 90-day exponential "
         "decay, so as patrol patterns shift, the model adapts.\n\n"
         "Strict chronological train, validation, and test split, with early stopping at "
         "iteration 54.\n\n"
         "The held-out test R-squared is 0.43 — visible right here on the dashboard. The "
         "train-to-test gap is just 0.07, well below the overfit threshold. Importantly, "
         "the test set actually performs slightly better than validation, which tells me "
         "the model generalizes — it isn't memorizing.\n\n"
         "What does the forecast tell us? Across the top 100 hotspots, roughly 743 illegal "
         "parking incidents will go unbooked in the next 24 hours, unless something "
         "changes.\n\n"
         "One honest caveat: the dataset captures bookings, not every event — a 1-week "
         "BTP pilot would validate the forecast in the wild."),

        ("3:35 – 3:55",
         "Scroll to the 'Top N illegal-parking hotspots' heatmap. Let Folium render. "
         "Briefly hover over the densest cluster (KR Market area).",
         "381 hotspots on the Bengaluru map. Densest cluster is Upparpet — KR Market. "
         "Every dot is clickable for station, junction, vehicle mix, and cost. Both the "
         "hotspot count and the number of patrols deployed are sliders in the sidebar — "
         "the system recomputes routes and ETAs live."),

        ("3:55 – 4:30",
         "Scroll to 'Tonight's optimized patrol plan'. Let the colour-coded route map render before speaking.",
         "So I converted the forecast into action.\n\n"
         "Weighted K-Means assigns hotspots across 5 patrols, balanced by expected "
         "catches. Within each patrol, nearest-neighbour TSP minimises travel time at "
         "20 km per hour, with 15 minutes of service time per stop.\n\n"
         "Note — 5 patrols is a demo scenario. The optimizer is parameterized, so in "
         "deployment, BTP would provide their actual nightly fleet size and the routes "
         "scale to match.\n\n"
         "For the demo I show a 5-hour evening window — 53 optimised stops across 5 "
         "patrols, roughly 96 expected catches. About 3.8× the current per-shift output. "
         "The window is parameterized — BTP can supply their actual roster, including "
         "the continuous towing patrols they already run."),

        ("4:30 – 5:00",
         "Scroll to 'Per-officer shift sheet — live tracking'. Pick a patrol from dropdown. "
         "Click 'Mark next stop done' 2-3 times so judges see the remaining route re-sequence "
         "and ETAs update live.",
         "Each patrol also gets a per-officer shift sheet — printable, downloadable, "
         "sendable on WhatsApp.\n\n"
         "Watch what happens as I mark stops complete. The remaining route re-sequences "
         "from the current position, ETAs update live, and if the replanned route runs "
         "past shift-end, the system flags which low-value stop to drop. That's real-time "
         "adaptive routing — working in v1 today.\n\n"
         "Below it sits the hotspot priority list — 381 zones, ranked, with the responsible "
         "police station, BTP junction code, dominant vehicle type, time-of-day share, "
         "and the actual address. Every prediction is interrogable. No black box."),

        ("5:00 – 5:10",
         "Slow-scroll back up to the top of the dashboard. Steady cursor.",
         "Park-Watch is the intelligence layer that closes Bengaluru's 12-hour enforcement "
         "blind spot. It runs on the data BTP already collects. No new sensors. No new "
         "cameras. And every officer keeps doing what they already do — just smarter.\n\n"
         "Park-Watch doesn't replace patrols. It makes them predictive instead of "
         "reactive — which is the actual gap the brief names.\n\n"
         "And the system gets better with use. Daily incremental updates, weekly full "
         "retrain on a rolling 6-month window, monthly hotspot re-clustering. Park-Watch "
         "bootstraps as the patrol patterns expand.\n\n"
         "v1 already routes adaptively per officer with live re-planning. v2 adds "
         "personalized learning over time.\n\n"
         "Thank you."),
    ]

    for ts, action, script in sections:
        _section_header(doc, ts, action)
        for para in script.split("\n\n"):
            _script_para(doc, para)

    # ── After recording
    doc.add_page_break()
    _heading(doc, "After you've recorded", size=14, after=6)

    steps = [
        ("1.", "Trim the start/end silence in QuickTime (Edit → Trim)."),
        ("2.", "Export — File → Export As → 1080p."),
        ("3.", "Upload to YouTube as Unlisted (do not post publicly until after judging closes)."),
        ("4.", "Paste the YouTube link into the HackerEarth 'Video Presentation' field."),
    ]
    for num, text in steps:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{num}  ")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = ORANGE
        r2 = p.add_run(text)
        r2.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.15)

    doc.add_paragraph()

    _heading(doc, "Tips for the live take", size=13, after=4)
    tips = [
        "Read the script before recording so the words feel natural — not robotic.",
        "Pause briefly at every full stop. Recording feels faster than it is.",
        "If you misspeak, pause for 2 seconds and restart the sentence — easier to edit later.",
        "Smile when you hit the multiplier line in the patrol section. Judges notice energy.",
        "The script targets ~3:00. If you naturally run to 3:15, that's fine. Aim for under 3:30 total.",
    ]
    for t in tips:
        p = doc.add_paragraph()
        r = p.add_run("•  " + t)
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.15)

    doc.save(OUT)
    print(f"Saved → {OUT}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
